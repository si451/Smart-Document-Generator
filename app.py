import streamlit as st
import fitz  # PyMuPDF
import docx
import os
import io
import tiktoken # To count tokens
from groq import Groq
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Load environment variables
load_dotenv()

# --- Helper Functions ---

def count_tokens(text, model="cl100k_base"):
    """Counts the number of tokens in a string of text."""
    encoding = tiktoken.get_encoding(model)
    return len(encoding.encode(text))

def extract_text_from_pdfs(pdf_files):
    """Extracts text from a list of uploaded PDF files."""
    st.write("Reading PDF files...")
    combined_text = ""
    for pdf_file in pdf_files:
        try:
            pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
            for page_num, page in enumerate(pdf_document):
                combined_text += page.get_text()
            pdf_document.close()
            st.write(f"  - Successfully read {pdf_file.name}")
        except Exception as e:
            st.error(f"Error processing {pdf_file.name}: {e}")
    return combined_text

def extract_text_from_docx(docx_file):
    """Extracts the full text content from a .docx file."""
    try:
        doc = docx.Document(io.BytesIO(docx_file.read()))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        docx_file.seek(0) # Reset file pointer for any subsequent use
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"Error reading the DOCX template: {e}")
        return None

def summarize_chunk(chunk, api_key):
    """Asks the LLM to summarize a single chunk of text."""
    client = Groq(api_key=api_key)
    prompt = f"""Summarize the key facts, figures, names, dates, and events from the following text for a claims report. Be concise and factual.

    **Text to Summarize:**
    {chunk}

    **Concise Summary:**
    """
    try:
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192", # Use a smaller model for fast summarization
            temperature=0.0,
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        st.warning(f"Could not summarize a chunk: {e}")
        return None

def generate_final_document(summarized_context, template_text, api_key):
    """Generates the final document from the summarized context."""
    client = Groq(api_key=api_key)
    prompt = f"""You are an expert claims adjuster. Rewrite the following Word document template, filling it in with information from the **Summarized Context** provided.

    Your final output must be **only the full, completed text of the report**. The structure and length should match the original template. Do not add extra headers or conversational text.

    --- 
    **Summarized Context:**
    {summarized_context}
    
    --- 
    **Word Document Template:**
    {template_text}
    """
    try:
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192", # Use the powerful model for the final generation
            temperature=0.1,
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        st.error(f"Final document generation failed: {e}")
        return None

def create_docx_from_text(text_content):
    """Creates a new .docx document from a string of text."""
    try:
        doc = docx.Document()
        for line in text_content.split('\n'):
            doc.add_paragraph(line)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        st.error(f"Error creating the final Word document: {e}")
        return None

# --- Streamlit UI ---

st.set_page_config(layout="centered", page_title="Smart Document Generator")
st.title("📄 Smart Document Generator")
st.markdown("Handles large documents by summarizing them first, then filling your template.")

GROQ_API_KEY = "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"

if not GROQ_API_KEY:
    st.warning("Groq API key not found.", icon="🔑")

st.subheader("1. Upload Your Template")
template_file = st.file_uploader("Upload a .docx template", type=["docx"])

st.subheader("2. Upload Information Reports")
report_files = st.file_uploader("Upload .pdf reports", type=["pdf"], accept_multiple_files=True)

st.divider()

if st.button("🚀 Generate Filled Document", type="primary", use_container_width=True, disabled=not (template_file and report_files and GROQ_API_KEY)):
    # Approx. token limit for the context window, leaving room for the prompt and response
    TOKEN_LIMIT = 6000 

    with st.spinner('Processing documents...'):
        pdf_text = extract_text_from_pdfs(report_files)
        template_text = extract_text_from_docx(template_file)

        if not pdf_text or not template_text:
            st.error("Failed to read files. Please check them and try again.")
            st.stop()

        total_tokens = count_tokens(pdf_text + template_text)
        st.write(f"Estimated total tokens: {total_tokens}")

        final_context = pdf_text

        # If tokens exceed limit, perform map-reduce summarization
        if total_tokens > TOKEN_LIMIT:
            st.info("Input is too large. Summarizing PDF content first...")
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=200)
            pdf_chunks = text_splitter.split_text(pdf_text)
            
            summaries = []
            summary_progress = st.progress(0)
            status_text = st.empty()
            for i, chunk in enumerate(pdf_chunks):
                status_text.text(f"Summarizing chunk {i+1}/{len(pdf_chunks)}...")
                summary = summarize_chunk(chunk, GROQ_API_KEY)
                if summary:
                    summaries.append(summary)
                summary_progress.progress((i + 1) / len(pdf_chunks))
            
            final_context = "\n\n".join(summaries)
            status_text.text("Summarization complete!")
            st.success(f"Condensed PDF content from {len(pdf_chunks)} chunks into one summary.")

        # Generate the final document
        st.info("Generating the final report...")
        final_document_text = generate_final_document(final_context, template_text, GROQ_API_KEY)

        if final_document_text:
            filled_document_stream = create_docx_from_text(final_document_text)
            if filled_document_stream:
                st.success("Document generated successfully!")
                st.download_button(
                    label="📥 Download Filled Document",
                    data=filled_document_stream,
                    file_name="Filled_Report_Final.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                with st.expander("View Generated Text"):
                    st.text_area("", final_document_text, height=300)
            else:
                st.error("Failed to create the final .docx file.")
        else:
            st.error("Failed to generate the document text from the LLM.")